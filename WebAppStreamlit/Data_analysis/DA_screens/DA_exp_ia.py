import streamlit as st
from fpdf import FPDF
from docx import Document
import os

def export_to_pdf(data, file_name):
    """Gera um relatório em PDF baseado nos dados fornecidos."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Cabeçalho
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(200, 10, txt="Relatório Técnico", ln=True, align='C')
    pdf.ln(10)

    # Seção 1: Informações Básicas
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 10, txt="1. Informações Básicas", ln=True)
    pdf.set_font("Arial", size=12)
    for key, value in data['basic_info'].items():
        pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)
    pdf.ln(10)

    # Seção 2: Métricas
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 10, txt="2. Métricas", ln=True)
    pdf.set_font("Arial", size=12)
    metrics = data.get('metrics', {})
    for key, value in metrics.items():
        pdf.cell(0, 10, txt=f"{key}: {value}", ln=True)
    pdf.ln(10)

    # Seção 3: Tabela de Resultados
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 10, txt="3. Tabela de Resultados", ln=True)
    pdf.set_font("Arial", size=10)

    if hasattr(data['results'], 'head'):
        header = list(data['results'].columns)
        rows = data['results'].head(20).values.tolist()

        # Configurações de largura das colunas
        column_widths = [30, 40, 20, 25, 25, 25, 30]  # Ajuste as larguras conforme necessário
        if len(header) > len(column_widths):  # Expande caso o número de colunas seja maior
            column_widths += [30] * (len(header) - len(column_widths))

        # Desenhar cabeçalho
        for i, col in enumerate(header):
            pdf.cell(column_widths[i], 10, col, border=1, align='C')
        pdf.ln()

        # Desenhar linhas
        for row in rows:
            for i, item in enumerate(row):
                pdf.cell(column_widths[i], 8, str(item), border=1, align='C')
            pdf.ln()
    else:
        pdf.cell(0, 10, txt="Sem dados disponíveis.", ln=True)
    pdf.ln(10)

    # Seção 4: Gráfico Interativo
    pdf.set_font("Arial", style="B", size=12)
    pdf.cell(200, 10, txt="4. Gráfico Interativo", ln=True)
    if "graph" in st.session_state:
        graph_image_path = "temp_graph.png"
        st.session_state["graph"].write_image(graph_image_path, format="png")
        pdf.image(graph_image_path, x=10, y=None, w=190)
        os.remove(graph_image_path)  # Remove o arquivo temporário
    else:
        pdf.cell(0, 10, txt="Gráfico não disponível", ln=True)
    pdf.ln(10)

    pdf.output(file_name)

def export_to_word(data, file_name):
    """Gera um relatório em Word baseado nos dados fornecidos."""
    doc = Document()

    # Cabeçalho
    doc.add_heading('Relatório Técnico', level=1)

    # Seção 1: Informações Básicas
    doc.add_heading('1. Informações Básicas', level=2)
    for key, value in data['basic_info'].items():
        doc.add_paragraph(f"{key}: {value}")

    # Seção 2: Métricas
    doc.add_heading('2. Métricas', level=2)
    metrics = data.get('metrics', {})
    for key, value in metrics.items():
        doc.add_paragraph(f"{key}: {value}")

    # Seção 3: Tabela de Resultados
    doc.add_heading('3. Tabela de Resultados', level=2)
    if hasattr(data['results'], 'head'):
        header = list(data['results'].columns)
        rows = data['results'].head(20).values.tolist()

        table = doc.add_table(rows=1, cols=len(header))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(header):
            hdr_cells[i].text = col

        for row in rows:
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
    else:
        doc.add_paragraph("Sem dados disponíveis.")

    # Seção 4: Gráfico Interativo
    doc.add_heading('4. Gráfico Interativo', level=2)
    if "graph" in st.session_state:
        graph_image_path = "temp_graph.png"
        st.session_state["graph"].write_image(graph_image_path, format="png")
        doc.add_picture(graph_image_path, width=5000000)
        os.remove(graph_image_path)  # Remove o arquivo temporário
    else:
        doc.add_paragraph("Gráfico não disponível.")

    doc.save(file_name)

def layout():
    st.title("Exportar Dados")
    export_option = st.selectbox("Escolha o formato de exportação:", ["PDF", "Word"])

    # Captura de informações diretamente do session_state
    data = {
        'basic_info': {
            'Unidade': st.session_state.get('unidade', 'Não informado'),
            'Empresa': st.session_state.get('empresa', 'Não informado'),
            'Data': st.session_state.get('data_coleta', 'Não informado'),
            'Equipe': st.session_state.get('equipe', 'Não informado'),
            'Horário Inicial': st.session_state.get('horario_inicial', 'Não informado'),
            'Horário Final': st.session_state.get('horario_final', 'Não informado'),
            'Diâmetro do Tubo': st.session_state.get('diametro_tubo', 'Não informado'),
            'Diâmetro Medido': st.session_state.get('diametro_medido', 'Não informado'),
            'Coordenadas': st.session_state.get('coordenadas', 'Não informado')
        },
        'metrics': st.session_state.get('processed_data', {}).get('metrics', {}),
        'results': st.session_state.get('processed_data', {}).get('table', [["Sem dados"]]),
        'collection_results': st.session_state.get('coleta_dados', {}).get('pontos_velocidades', [["Sem dados"]]),
    }

    if st.button("Exportar"):
        # Geração do relatório
        if export_option == "PDF":
            file_name = "relatorio.pdf"
            export_to_pdf(data, file_name)
            mime_type = "application/pdf"
        elif export_option == "Word":
            file_name = "relatorio.docx"
            export_to_word(data, file_name)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # Botão de download
        with open(file_name, "rb") as f:
            st.download_button(
                label=f"Baixar Relatório ({export_option})",
                data=f.read(),
                file_name=file_name,
                mime=mime_type
            )

        os.remove(file_name)

if __name__ == "__main__":
    layout()
